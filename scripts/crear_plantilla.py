"""
Genera la plantilla base plantilla_ficha.docx en la carpeta plantillas/.
Ejecútalo una vez; luego abre el .docx en Word y personaliza fuentes,
colores y márgenes a tu gusto sin cambiar los marcadores {{ campo }}.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECCIONES = [
    ("Resumen factual",                  "resumen"),
    ("Elementos prioritarios extraídos", "elementos_prioritarios"),
    ("Términos clave identificados",     "terminos_clave"),
    ("Datos relevantes adicionales",     "datos_adicionales"),
    ("Elementos excluidos",              "elementos_excluidos"),
]

METADATOS = [
    ("Título original",  "titulo_original"),
    ("Tipo de documento","tipo_documento"),
    ("Autor(es)",        "autores"),
    ("Fecha",            "fecha"),
    ("Fuente / origen",  "fuente"),
]


def crear_plantilla(output_path: Path):
    doc = Document()

    # Márgenes de página
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── Título ───────────────────────────────────────────────────────────────
    doc.add_paragraph("{{ titulo }}", style="Heading 1")

    # ── Metadatos ────────────────────────────────────────────────────────────
    doc.add_paragraph("Datos del documento fuente", style="Heading 2")
    for label, campo in METADATOS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run("{{ " + campo + " }}")

    # ── Secciones de contenido ───────────────────────────────────────────────
    for titulo_seccion, campo in SECCIONES:
        doc.add_paragraph(titulo_seccion, style="Heading 2")
        # Bucle a nivel de párrafo: {%p %} indica a docxtpl que este párrafo
        # es una etiqueta de bloque, no contenido visible.
        doc.add_paragraph("{%p for para in " + campo + "_items %}")
        doc.add_paragraph("{{ para }}")
        doc.add_paragraph("{%p endfor %}")

    doc.save(output_path)
    print(f"Plantilla generada en: {output_path}")
    print("Puedes abrirla en Word y personalizar estilos sin tocar los marcadores {{ }}.")


if __name__ == "__main__":
    base = Path(__file__).parent.parent
    output = base / "plantillas" / "plantilla_ficha.docx"
    crear_plantilla(output)
